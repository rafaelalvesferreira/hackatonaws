import os
import logging
from typing import List, Dict
import PyPDF2
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extrai texto de arquivo PDF"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Erro ao processar PDF {file_path}: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extrai texto de arquivo Word"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Erro ao processar DOCX {file_path}: {str(e)}")
            return ""
    
    def process_document(self, file_path: str) -> List[LangchainDocument]:
        """Processa um documento e retorna chunks"""
        file_name = os.path.basename(file_path)
        file_extension = os.path.splitext(file_path)[1].lower()
        
        logger.info(f"Processando documento: {file_name}")
        
        # Extrai texto baseado na extensão
        if file_extension == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            text = self.extract_text_from_docx(file_path)
        else:
            logger.warning(f"Tipo de arquivo não suportado: {file_extension}")
            return []
        
        if not text.strip():
            logger.warning(f"Nenhum texto extraído de {file_name}")
            return []
        
        # Divide o texto em chunks
        chunks = self.text_splitter.split_text(text)
        
        # Cria documentos Langchain com metadados
        documents = []
        for i, chunk in enumerate(chunks):
            doc = LangchainDocument(
                page_content=chunk,
                metadata={
                    'source': file_name,
                    'chunk_id': i,
                    'file_path': file_path,
                    'file_type': file_extension
                }
            )
            documents.append(doc)
        
        logger.info(f"Documento {file_name} processado em {len(documents)} chunks")
        return documents
    
    def process_documents_directory(self, directory_path: str) -> List[LangchainDocument]:
        """Processa todos os documentos em um diretório"""
        all_documents = []
        supported_extensions = ['.pdf', '.docx', '.doc']
        
        if not os.path.exists(directory_path):
            logger.error(f"Diretório não encontrado: {directory_path}")
            return all_documents
        
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            file_extension = os.path.splitext(filename)[1].lower()
            
            if file_extension in supported_extensions and os.path.isfile(file_path):
                print("File Path:", file_path)
                documents = self.process_document(file_path)
                all_documents.extend(documents)
        
        logger.info(f"Total de {len(all_documents)} chunks processados de {directory_path}")
        return all_documents
